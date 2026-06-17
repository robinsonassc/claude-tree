#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""dirtree — generate a visual directory tree and optionally save it to a file.

Cross-platform (Windows/macOS/Linux). Pure standard library, except that the
optional ``--format docx`` output requires ``python-docx`` (``pip install python-docx``).

Examples
--------
    python dirtree.py .                         # print tree of current dir
    python dirtree.py /path/to/proj -o tree.txt # save plain-text tree
    python dirtree.py . -f md  -o STRUCTURE.md   # markdown (fenced) tree
    python dirtree.py . -f docx -o Index.docx    # styled Word document
    python dirtree.py . -d --depth 2             # dirs only, 2 levels deep
"""
from __future__ import annotations

import argparse
import fnmatch
import os
import sys
from datetime import datetime
from pathlib import Path

PIPE, TEE, ELBOW, SPACE = "│   ", "├── ", "└── ", "    "

DEFAULT_IGNORES = [
    ".git", ".svn", ".hg", "node_modules", "__pycache__", ".venv", "venv",
    ".next", "dist", "build", ".DS_Store", "desktop.ini", "Thumbs.db",
    "~$*", "*.pyc",
]


def _ignored(name: str, patterns: list[str]) -> bool:
    return any(fnmatch.fnmatch(name, pat) for pat in patterns)


def _entries(path: Path, patterns: list[str], dir_only: bool) -> list[Path]:
    try:
        children = list(path.iterdir())
    except (PermissionError, OSError):
        return []
    kept = []
    for child in children:
        if _ignored(child.name, patterns):
            continue
        if dir_only and child.is_file():
            continue
        kept.append(child)
    # directories first, then files, each alphabetical (case-insensitive)
    return sorted(kept, key=lambda c: (c.is_file(), c.name.lower()))


def build_tree(root: Path, patterns: list[str], dir_only: bool, max_depth: int | None):
    """Return (lines, dir_count, file_count) for the tree under root."""
    lines: list[str] = []
    counts = {"dirs": 0, "files": 0}

    def walk(path: Path, prefix: str, depth: int) -> None:
        if max_depth is not None and depth >= max_depth:
            return
        items = _entries(path, patterns, dir_only)
        for index, child in enumerate(items):
            last = index == len(items) - 1
            connector = ELBOW if last else TEE
            if child.is_dir():
                counts["dirs"] += 1
                lines.append(prefix + connector + child.name + "/")
                walk(child, prefix + (SPACE if last else PIPE), depth + 1)
            else:
                counts["files"] += 1
                lines.append(prefix + connector + child.name)

    walk(root, "", 0)
    return lines, counts["dirs"], counts["files"]


def render_text(title: str, root_name: str, lines: list[str], summary: str, stamp: str) -> str:
    head = [title, "Generated: " + stamp, "", root_name + "/"]
    return "\n".join(head + lines + ["", summary]) + "\n"


def render_markdown(title: str, root_name: str, lines: list[str], summary: str, stamp: str) -> str:
    body = "\n".join([root_name + "/"] + lines)
    return (
        f"# {title}\n\n"
        f"_Generated: {stamp}_\n\n"
        f"```text\n{body}\n```\n\n"
        f"{summary}\n"
    )


def write_docx(out: Path, title: str, root_name: str, lines: list[str], summary: str, stamp: str) -> None:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        sys.exit("--format docx requires python-docx. Install with: pip install python-docx")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    h = doc.add_paragraph()
    run = h.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 0, 0)

    meta = doc.add_paragraph()
    mrun = meta.add_run("Generated: " + stamp)
    mrun.italic = True
    mrun.font.size = Pt(10)
    mrun.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    body = doc.add_paragraph()
    brun = body.add_run("\n".join([root_name + "/"] + lines))
    brun.font.name = "Consolas"   # monospace keeps the connectors aligned
    brun.font.size = Pt(9)

    s = doc.add_paragraph()
    s.add_run(summary).italic = True
    doc.save(str(out))


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Generate a directory tree; optionally save to a file.")
    parser.add_argument("root", nargs="?", default=".", help="Directory to map (default: current directory)")
    parser.add_argument("-o", "--output", help="Write the tree to this file (otherwise prints to stdout)")
    parser.add_argument("-f", "--format", choices=["txt", "md", "docx"], default="txt",
                        help="Output format (default: txt). docx needs python-docx.")
    parser.add_argument("-d", "--dir-only", action="store_true", help="Show directories only")
    parser.add_argument("--depth", type=int, default=None, help="Limit tree depth (levels below root)")
    parser.add_argument("--ignore", default="", help="Comma-separated extra names/globs to ignore")
    parser.add_argument("--no-default-ignores", action="store_true",
                        help="Do not apply the built-in ignore list")
    parser.add_argument("--no-counts", action="store_true", help="Omit the trailing dir/file summary")
    parser.add_argument("--title", default=None, help="Override the document title")
    args = parser.parse_args(argv)

    root = Path(args.root).expanduser().resolve()
    if not root.is_dir():
        sys.exit(f"Not a directory: {root}")

    patterns = [] if args.no_default_ignores else list(DEFAULT_IGNORES)
    patterns += [p.strip() for p in args.ignore.split(",") if p.strip()]
    # never recurse into our own output file
    if args.output:
        patterns.append(Path(args.output).name)

    lines, ndirs, nfiles = build_tree(root, patterns, args.dir_only, args.depth)
    title = args.title or f"{root.name} — Folder Index"
    stamp = datetime.now().strftime("%d %B %Y, %I:%M %p")
    summary = "" if args.no_counts else f"{ndirs} directories, {nfiles} files"

    if args.format == "docx":
        if not args.output:
            sys.exit("--format docx requires -o/--output (a .docx path)")
        write_docx(Path(args.output), title, root.name, lines, summary, stamp)
        print(f"Wrote {args.output} ({ndirs} dirs, {nfiles} files)")
        return 0

    if args.format == "md":
        text = render_markdown(title, root.name, lines, summary, stamp)
    else:
        text = render_text(title, root.name, lines, summary, stamp)

    if args.output:
        Path(args.output).write_text(text, encoding="utf-8")
        print(f"Wrote {args.output} ({ndirs} dirs, {nfiles} files)")
    else:
        # ensure box-drawing characters survive a non-UTF8 console (e.g. Windows cp1252)
        if hasattr(sys.stdout, "buffer"):
            sys.stdout.buffer.write(text.encode("utf-8"))
        else:
            print(text)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
